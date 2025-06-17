from docx import Document
from docx.shared import Pt

def save_contract_to_docx(deal, filename):
    """Сохраняет договор в формате .docx."""

    doc = Document()
    doc.add_heading('Договор купли-продажи квартиры', level=1)


    # Устанавливаем стиль шрифта
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(10)

    contract_text = deal.generate_contract()

    for line in contract_text.split("\n"):
        doc.add_paragraph(line)

    doc.save(filename)
    print(f"Договор сохранён в файл: {filename}")
